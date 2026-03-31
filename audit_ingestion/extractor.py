"""
audit_ingestion_v04.2/audit_ingestion/extractor.py
Page-aware extraction engine with fast/deep lane split.

Fast Review (extract_fast):
  - pdfplumber text only (no table scan)
  - shared PyPDF2 reader for weak pages
  - no OCR, no vision
  - escalate globally if total extraction is weak

Deep Extraction (extract_deep):
  - pdfplumber text + targeted table scan
  - shared PyPDF2 reader
  - OCR on weak pages (top 6 max)
  - vision on critical pages (top 2 max)

Shared PDF handles — never reopen per page.
"""
from __future__ import annotations
import hashlib
import io
import logging
from pathlib import Path
from typing import Optional
from .models import ParsedDocument, ParsedPage, ParsedTable

logger = logging.getLogger(__name__)

# ── Limits ────────────────────────────────────────────────────────────────────
MAX_PDF_PAGES_FAST        = 40
MAX_PDF_PAGES_TABLE_SCAN  = 12
MAX_OCR_PAGES             = 6
MAX_VISION_PAGES          = 2
MAX_RELEVANT_PAGES        = 10
MAX_CONTEXT_CHARS         = 20000

# ── Thresholds ────────────────────────────────────────────────────────────────
MIN_CHARS_ACCEPTABLE  = 350
MIN_CHARS_WEAK        = 150
MIN_CHARS_CRITICAL    = 60


# ── Cache (in-process, keyed by file hash) ────────────────────────────────────
_extraction_cache: dict[str, ParsedDocument] = {}
_ocr_page_cache:  dict[str, str] = {}       # key: file_hash:page_index
_image_page_cache: dict[str, bytes] = {}    # key: file_hash:page_index:dpi


def _file_hash(path: str) -> str:
    h = hashlib.md5()
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(65536), b""):
            h.update(chunk)
    return h.hexdigest()


# ── pdfplumber extraction ─────────────────────────────────────────────────────

def _pdfplumber_extract(
    path: str,
    max_pages: int,
    scan_tables: bool,
    table_pages_limit: int = MAX_PDF_PAGES_TABLE_SCAN,
) -> tuple[list[ParsedPage], list[ParsedTable], int]:
    try:
        import pdfplumber
    except ImportError:
        return [], [], 0

    pages: list[ParsedPage] = []
    tables: list[ParsedTable] = []

    try:
        with pdfplumber.open(path) as pdf:
            page_count = len(pdf.pages)
            limit = min(page_count, max_pages)

            for i, page in enumerate(pdf.pages[:limit]):
                text = page.extract_text() or ""
                conf = min(1.0, len(text) / 500) if text else 0.0
                pages.append(ParsedPage(
                    page_number=i + 1,
                    text=text,
                    char_count=len(text),
                    extractor="pdfplumber",
                    confidence=conf,
                ))

                if scan_tables and i < table_pages_limit:
                    for ti, tbl in enumerate(page.extract_tables() or []):
                        if not tbl or len(tbl) < 2:
                            continue
                        headers = tbl[0]
                        valid_h = sum(1 for h in headers if h and str(h).strip())
                        if valid_h >= 2:
                            clean_h = [str(h).strip() if h else f"Col_{j}"
                                       for j, h in enumerate(headers)]
                            rows = [
                                {clean_h[ci] if ci < len(clean_h) else f"Col_{ci}": cell
                                 for ci, cell in enumerate(row)}
                                for row in tbl[1:]
                            ]
                            tables.append(ParsedTable(
                                page_number=i + 1, table_index=ti,
                                headers=clean_h, rows=rows,
                                row_count=len(rows), extractor="pdfplumber",
                            ))
        return pages, tables, page_count
    except Exception as e:
        logger.warning(f"pdfplumber failed: {e}")
        return [], [], 0


# ── PyPDF2 — shared reader, no reopen per page ────────────────────────────────

def _load_pypdf2_reader(path: str):
    try:
        from PyPDF2 import PdfReader
        with open(path, "rb") as f:
            return PdfReader(io.BytesIO(f.read()))
    except Exception:
        return None


def _pypdf2_page_text(reader, page_index: int) -> str:
    try:
        if reader and page_index < len(reader.pages):
            return reader.pages[page_index].extract_text() or ""
    except Exception:
        pass
    return ""


# ── extractous full-document pass ─────────────────────────────────────────────

def _extractous_full(path: str) -> dict[int, str]:
    try:
        from extractous import Extractor, TesseractOcrConfig
        extractor = (
            Extractor()
            .set_extract_string_max_length(2_000_000)
            .set_ocr_config(TesseractOcrConfig().set_language("eng"))
        )
        text, _ = extractor.extract_file_to_string(str(path))
        if not text or not text.strip():
            return {}
        page_splits = text.split("\x0c")
        return {i + 1: t for i, t in enumerate(page_splits) if t.strip()}
    except Exception as e:
        logger.debug(f"extractous failed: {e}")
        return {}


# ── OCR — shared fitz doc, no reopen per page ────────────────────────────────

def _load_fitz_doc(path: str):
    try:
        import fitz
        return fitz.open(path)
    except Exception:
        return None


def _ocr_page(fitz_doc, page_index: int, dpi: int = 250,
              cache_key_prefix: str = "") -> str:
    """
    OCR a single page. Uses in-process cache keyed by file+page.
    Strategy: psm 6 (uniform block) first — better for tables and forms.
    Falls back to psm 11 (sparse text) if psm 6 yields less content.
    Uses 300 DPI for best table extraction quality.
    """
    ocr_dpi = max(dpi, 300)  # Always use at least 300 DPI for OCR
    ck = f"{cache_key_prefix}:{page_index}:{ocr_dpi}"
    if ck in _ocr_page_cache:
        return _ocr_page_cache[ck]
    try:
        import pytesseract
        from PIL import Image
        page = fitz_doc[page_index]
        pix = page.get_pixmap(dpi=ocr_dpi, alpha=False)
        mode = "RGB" if pix.n < 4 else "RGBA"
        img = Image.frombytes(mode, [pix.width, pix.height], pix.samples)

        # psm 6 = uniform block — better for table/form layouts (invoices, leases, schedules)
        text_psm6 = pytesseract.image_to_string(img, lang="eng", config="--oem 1 --psm 6")
        # psm 11 = sparse text — better for loose/mixed layouts
        text_psm11 = pytesseract.image_to_string(img, lang="eng", config="--oem 1 --psm 11")

        # Pick whichever produced more content
        text = text_psm6 if len(text_psm6.strip()) >= len(text_psm11.strip()) else text_psm11

        if ck:
            _ocr_page_cache[ck] = text
        return text
    except Exception as e:
        logger.debug(f"OCR page {page_index + 1} failed: {e}")
        return ""


# ── Vision — render page images ───────────────────────────────────────────────

def _render_page_images(fitz_doc, page_indices: list[int],
                        dpi: int = 150, file_hash: str | None = None) -> list[bytes]:
    images: list[bytes] = []
    try:
        from PIL import Image
        import io as _io
        for idx in page_indices:
            if idx >= len(fitz_doc):
                continue

            cache_key = f"{file_hash}:{idx}:{dpi}" if file_hash else None
            if cache_key and cache_key in _image_page_cache:
                images.append(_image_page_cache[cache_key])
                continue

            page = fitz_doc[idx]
            pix = page.get_pixmap(dpi=dpi, alpha=False)
            mode = "RGB" if pix.n < 4 else "RGBA"
            img = Image.frombytes(mode, [pix.width, pix.height], pix.samples)
            buf = _io.BytesIO()
            img.save(buf, format="JPEG", quality=85)
            data = buf.getvalue()
            if cache_key:
                _image_page_cache[cache_key] = data
            images.append(data)
    except Exception as e:
        logger.warning(f"Page render failed: {e}")
    return images


def render_page_image_cached(path: str, page_index: int, dpi: int = 200) -> bytes:
    """
    Render a single PDF page to JPEG bytes.
    Cached by file hash + page index + dpi.
    Used by rescue path in router — independent of shared fitz handle.
    """
    file_hash = _file_hash(path)
    cache_key = f"{file_hash}:{page_index}:{dpi}"
    if cache_key in _image_page_cache:
        return _image_page_cache[cache_key]

    try:
        import fitz
        import io as _io
        from PIL import Image
    except ImportError:
        return b""

    doc = None
    try:
        doc = fitz.open(path)
        page = doc[page_index]
        pix = page.get_pixmap(dpi=dpi, alpha=False)
        mode = "RGB" if pix.n < 4 else "RGBA"
        img = Image.frombytes(mode, [pix.width, pix.height], pix.samples)
        buf = _io.BytesIO()
        img.save(buf, format="JPEG", quality=90)
        data = buf.getvalue()
        _image_page_cache[cache_key] = data
        return data
    except Exception:
        return b""
    finally:
        try:
            if doc:
                doc.close()
        except Exception:
            pass


def _vision_weak_pages(
    fitz_doc,
    weak_indices: list[int],
    provider,
    file_hash: str | None = None,
) -> dict[int, str]:
    if not weak_indices or not provider:
        return {}
    if not hasattr(provider, "extract_text_from_page_images"):
        return {}

    cap = min(len(weak_indices), MAX_VISION_PAGES)
    target_indices = weak_indices[:cap]
    images = _render_page_images(fitz_doc, target_indices, file_hash=file_hash)
    if not images:
        return {}

    try:
        prompt = (
            "Extract ALL text from these document pages faithfully. "
            "Preserve all numbers, dates, names, amounts, and terms exactly as written. "
            "Separate pages with '--- PAGE BREAK ---'."
        )
        combined = provider.extract_text_from_page_images(images=images, prompt=prompt)
        page_texts = combined.split("--- PAGE BREAK ---")
        return {
            target_indices[i]: text.strip()
            for i, text in enumerate(page_texts)
            if i < len(target_indices) and text.strip()
        }
    except Exception as e:
        logger.warning(f"Vision extraction failed: {e}")
        return {}


# ── Non-PDF direct extraction ─────────────────────────────────────────────────

def _extract_direct(path: str) -> ParsedDocument:
    p = Path(path)
    ext = p.suffix.lower()
    tables: list[ParsedTable] = []
    text = ""

    fh = _file_hash(path)
    try:
        if ext in (".csv", ".tsv"):
            import pandas as pd
            df = pd.read_csv(path)
            text = df.to_string(index=False)
            tables = [ParsedTable(
                page_number=1,
                headers=df.columns.tolist(),
                rows=df.head(200).to_dict("records"),
                row_count=len(df),
                extractor="direct",
            )]

        elif ext in (".xlsx", ".xls"):
            import pandas as pd
            df = pd.read_excel(path)
            text = df.to_string(index=False)
            tables = [ParsedTable(
                page_number=1,
                headers=df.columns.tolist(),
                rows=df.head(200).to_dict("records"),
                row_count=len(df),
                extractor="direct",
            )]

        elif ext == ".txt":
            text = p.read_text(encoding="utf-8", errors="replace")

        elif ext == ".docx":
            from docx import Document
            doc = Document(path)
            text = "\n".join(para.text for para in doc.paragraphs if para.text.strip())

        else:
            return ParsedDocument(
                source_file=p.name,
                errors=[f"Unsupported file type: {ext}"],
            )

        pg = ParsedPage(
            page_number=1, text=text, char_count=len(text),
            extractor="direct", confidence=1.0 if text else 0.0,
        )
        return ParsedDocument(
            source_file=p.name,
            file_hash=fh,
            mime_type=f"text/{ext.lstrip('.')}",
            full_text=text,
            page_count=1,
            pages=[pg],
            tables=[t.model_dump() for t in tables],
            extraction_chain=["direct"],
            primary_extractor="direct",
            confidence=1.0 if text else 0.0,
        )

    except Exception as e:
        return ParsedDocument(source_file=p.name, errors=[str(e)])


# ── Assemble ParsedDocument from pages ───────────────────────────────────────

def _assemble(
    source_file: str,
    file_hash: str,
    pages: list[ParsedPage],
    tables: list[ParsedTable],
    page_count: int,
    extraction_chain: list[str],
    weak_pages: list[int],
    ocr_pages: list[int],
    vision_pages: list[int],
    warnings: list[str],
    errors: list[str],
) -> ParsedDocument:
    full_text = "\n\n".join(
        f"[Page {pg.page_number}]\n{pg.text}"
        for pg in pages if pg.text.strip()
    )
    total_chars = sum(pg.char_count for pg in pages)
    n = max(page_count, len(pages), 1)
    confidence = min(1.0, (total_chars / n) / 500)

    # Compute primary_extractor from actual final page outputs
    extractor_counts: dict[str, int] = {}
    for pg in pages:
        extractor_counts[pg.extractor] = extractor_counts.get(pg.extractor, 0) + pg.char_count
    primary = max(extractor_counts, key=extractor_counts.get) if extractor_counts else "none"

    weak_final = [pg.page_number for pg in pages if pg.char_count < MIN_CHARS_ACCEPTABLE]
    if weak_final:
        warnings.append(f"{len(weak_final)} page(s) below acceptable threshold")

    return ParsedDocument(
        source_file=source_file,
        file_hash=file_hash,
        mime_type="application/pdf",
        full_text=full_text,
        page_count=n,
        pages=pages,
        tables=[t.model_dump() if hasattr(t, "model_dump") else t for t in tables],
        extraction_chain=list(dict.fromkeys(extraction_chain)),
        primary_extractor=primary,
        confidence=confidence,
        weak_pages=weak_final,
        ocr_pages=ocr_pages,
        vision_pages=vision_pages,
        warnings=warnings,
        errors=errors,
    )


# ── Fast lane ─────────────────────────────────────────────────────────────────

def extract_fast(path: str, provider=None, ocr_semaphore=None) -> ParsedDocument:
    """
    Fast Review lane:
    - pdfplumber text (no table scan)
    - shared PyPDF2 for weak pages
    - no OCR, no vision
    - one extractous pass if globally weak
    """
    cache_key = _file_hash(path)
    cache_hit = _extraction_cache.get(f"fast:{cache_key}")
    if cache_hit:
        return cache_hit

    p = Path(path)
    extraction_chain: list[str] = []
    warnings: list[str] = []
    errors: list[str] = []

    # pdfplumber — text only, no table scan
    pages, tables, page_count = _pdfplumber_extract(
        path, MAX_PDF_PAGES_FAST, scan_tables=False
    )
    if pages:
        extraction_chain.append("pdfplumber")

    # Shared PyPDF2 reader for weak pages
    pypdf2_reader = _load_pypdf2_reader(path)
    if pypdf2_reader:
        pypdf2_used = False
        for i, pg in enumerate(pages):
            if pg.char_count < MIN_CHARS_WEAK:
                t = _pypdf2_page_text(pypdf2_reader, i)
                if len(t) > pg.char_count:
                    pages[i] = ParsedPage(
                        page_number=pg.page_number, text=t, char_count=len(t),
                        extractor="pypdf2", confidence=min(1.0, len(t) / 500),
                    )
                    pypdf2_used = True
        if pypdf2_used:
            extraction_chain.append("pypdf2")

    # Global weak check — if still very low overall, try extractous once
    total_chars = sum(pg.char_count for pg in pages)
    n = max(page_count, len(pages), 1)
    if total_chars / n < MIN_CHARS_WEAK:
        ext_pages = _extractous_full(path)
        if ext_pages:
            extraction_chain.append("extractous")
            for pn, t in ext_pages.items():
                idx = pn - 1
                if idx < len(pages) and len(t) > pages[idx].char_count:
                    pages[idx] = ParsedPage(
                        page_number=pn, text=t, char_count=len(t),
                        extractor="extractous", confidence=min(1.0, len(t) / 500),
                    )
            if not pages:
                pages = [
                    ParsedPage(page_number=pn, text=t, char_count=len(t),
                               extractor="extractous", confidence=min(1.0, len(t) / 500))
                    for pn, t in ext_pages.items()
                ]
                page_count = len(pages)

    # Smart escalation — if fast mode is still critically weak, run OCR on worst pages
    total_chars_fast = sum(pg.char_count for pg in pages)
    n_pages = max(page_count, len(pages), 1)
    avg_cpp = total_chars_fast / n_pages
    ocr_pages_fast: list[int] = []
    fitz_doc_fast = None  # initialized here so cleanup pass can safely reference it

    if avg_cpp < MIN_CHARS_CRITICAL and pages:
        logger.info(f"Fast mode critically weak ({avg_cpp:.0f} chars/page) — escalating weak pages to OCR")
        fitz_doc_fast = _load_fitz_doc(path)
        if fitz_doc_fast:
            # Always OCR first 2 pages (most likely to contain key terms/header info)
            # Then fill remaining slots with lowest-char pages
            priority_indices = [i for i in [0, 1] if i < len(pages)
                                and pages[i].char_count < MIN_CHARS_CRITICAL]
            remaining_critical = sorted(
                [i for i, pg in enumerate(pages)
                 if pg.char_count < MIN_CHARS_CRITICAL and i not in priority_indices],
                key=lambda i: pages[i].char_count,
            )
            critical_indices = (priority_indices + remaining_critical)[:MAX_OCR_PAGES]

            ocr_used = False
            fhash = _file_hash(path)
            for i in critical_indices:
                if ocr_semaphore:
                    with ocr_semaphore:
                        t = _ocr_page(fitz_doc_fast, i, cache_key_prefix=fhash)
                else:
                    t = _ocr_page(fitz_doc_fast, i, cache_key_prefix=fhash)
                if len(t.strip()) > pages[i].char_count:
                    pages[i] = ParsedPage(
                        page_number=pages[i].page_number, text=t,
                        char_count=len(t), extractor="ocr",
                        confidence=0.75, image_used=True,
                    )
                    ocr_pages_fast.append(pages[i].page_number)
                    ocr_used = True
            if ocr_used:
                extraction_chain.append("ocr_escalated")

    # Cleanup pass — OCR any pages still critical after the main escalation
    # This catches pages that were bumped by the MAX_OCR_PAGES cap
    still_critical = [
        i for i, pg in enumerate(pages)
        if pg.char_count < MIN_CHARS_CRITICAL and pg.page_number not in ocr_pages_fast
    ]
    if still_critical and fitz_doc_fast:
        fhash2 = _file_hash(path)
        for i in still_critical[:3]:  # Cap cleanup at 3 additional pages
            t = _ocr_page(fitz_doc_fast, i, cache_key_prefix=fhash2)
            if len(t.strip()) > pages[i].char_count:
                pages[i] = ParsedPage(
                    page_number=pages[i].page_number, text=t,
                    char_count=len(t), extractor="ocr",
                    confidence=0.75, image_used=True,
                )
                ocr_pages_fast.append(pages[i].page_number)
        if still_critical:
            extraction_chain.append("ocr_cleanup")

    result = _assemble(
        p.name, cache_key, pages, tables, page_count, extraction_chain,
        [], ocr_pages_fast, [], warnings, errors,
    )
    _extraction_cache[f"fast:{cache_key}"] = result
    return result


# ── Deep lane ─────────────────────────────────────────────────────────────────

def extract_deep(path: str, provider=None, ocr_semaphore=None) -> ParsedDocument:
    """
    Deep Extraction lane:
    - pdfplumber text + targeted table scan
    - shared PyPDF2 reader
    - extractous if globally weak
    - OCR top-N weak pages (shared fitz doc)
    - vision top-N critical pages (shared fitz doc)
    """
    cache_key = _file_hash(path)
    cache_hit = _extraction_cache.get(f"deep:{cache_key}")
    if cache_hit:
        return cache_hit

    p = Path(path)
    extraction_chain: list[str] = []
    warnings: list[str] = []
    errors: list[str] = []
    ocr_pages: list[int] = []
    vision_pages: list[int] = []

    # pdfplumber with table scan
    pages, tables, page_count = _pdfplumber_extract(
        path, MAX_PDF_PAGES_FAST, scan_tables=True,
        table_pages_limit=MAX_PDF_PAGES_TABLE_SCAN,
    )
    if pages:
        extraction_chain.append("pdfplumber")

    # Shared PyPDF2 reader
    pypdf2_reader = _load_pypdf2_reader(path)
    if pypdf2_reader:
        pypdf2_used = False
        for i, pg in enumerate(pages):
            if pg.char_count < MIN_CHARS_WEAK:
                t = _pypdf2_page_text(pypdf2_reader, i)
                if len(t) > pg.char_count:
                    pages[i] = ParsedPage(
                        page_number=pg.page_number, text=t, char_count=len(t),
                        extractor="pypdf2", confidence=min(1.0, len(t) / 500),
                    )
                    pypdf2_used = True
        if pypdf2_used:
            extraction_chain.append("pypdf2")

    # extractous if globally weak
    total_chars = sum(pg.char_count for pg in pages)
    n = max(page_count, len(pages), 1)
    if total_chars / n < MIN_CHARS_WEAK:
        ext_pages = _extractous_full(path)
        if ext_pages:
            extraction_chain.append("extractous")
            for pn, t in ext_pages.items():
                idx = pn - 1
                if idx < len(pages) and len(t) > pages[idx].char_count:
                    pages[idx] = ParsedPage(
                        page_number=pn, text=t, char_count=len(t),
                        extractor="extractous", confidence=min(1.0, len(t) / 500),
                    )
            if not pages:
                pages = [
                    ParsedPage(page_number=pn, text=t, char_count=len(t),
                               extractor="extractous", confidence=min(1.0, len(t) / 500))
                    for pn, t in ext_pages.items()
                ]
                page_count = len(pages)

    # OCR — shared fitz doc, weak pages
    # Always prioritize first 2 pages (most likely to contain key terms/header)
    # then fill remaining slots by lowest char count
    _weak_all = [i for i, pg in enumerate(pages) if pg.char_count < MIN_CHARS_WEAK]
    _priority  = [i for i in [0, 1] if i in _weak_all]
    _remaining = sorted(
        [i for i in _weak_all if i not in _priority],
        key=lambda i: pages[i].char_count,
    )
    weak_indices = (_priority + _remaining)[:MAX_OCR_PAGES]

    if weak_indices:
        fitz_doc = _load_fitz_doc(path)
        if fitz_doc:
            ocr_used = False
            fhash_deep = _file_hash(path)
            for i in weak_indices:
                if ocr_semaphore:
                    with ocr_semaphore:
                        t = _ocr_page(fitz_doc, i, cache_key_prefix=fhash_deep)
                else:
                    t = _ocr_page(fitz_doc, i, cache_key_prefix=fhash_deep)
                if len(t.strip()) > pages[i].char_count:
                    pages[i] = ParsedPage(
                        page_number=pages[i].page_number, text=t,
                        char_count=len(t), extractor="ocr",
                        confidence=0.75, image_used=True,
                    )
                    ocr_pages.append(pages[i].page_number)
                    ocr_used = True

            if ocr_used:
                extraction_chain.append("ocr")

            # Cleanup pass — OCR any pages still critical after the main pass
            still_critical = [
                i for i, pg in enumerate(pages)
                if pg.char_count < MIN_CHARS_CRITICAL
                and pg.page_number not in ocr_pages
            ]
            for i in still_critical[:3]:
                if ocr_semaphore:
                    with ocr_semaphore:
                        t = _ocr_page(fitz_doc, i, cache_key_prefix=fhash_deep)
                else:
                    t = _ocr_page(fitz_doc, i, cache_key_prefix=fhash_deep)
                if len(t.strip()) > pages[i].char_count:
                    pages[i] = ParsedPage(
                        page_number=pages[i].page_number, text=t,
                        char_count=len(t), extractor="ocr",
                        confidence=0.75, image_used=True,
                    )
                    ocr_pages.append(pages[i].page_number)
            if still_critical:
                extraction_chain.append("ocr_cleanup")

            # Vision — still critical after OCR, top pages only
            critical_indices = sorted(
                [i for i, pg in enumerate(pages) if pg.char_count < MIN_CHARS_CRITICAL],
                key=lambda i: pages[i].char_count,
            )[:MAX_VISION_PAGES]

            if critical_indices and provider is not None:
                vision_results = _vision_weak_pages(fitz_doc, critical_indices, provider, file_hash=fhash_deep)
                if vision_results:
                    extraction_chain.append("vision")
                    for i, t in vision_results.items():
                        if i < len(pages) and len(t) > pages[i].char_count:
                            pages[i] = ParsedPage(
                                page_number=pages[i].page_number, text=t,
                                char_count=len(t), extractor="vision",
                                confidence=0.85, image_used=True,
                            )
                            vision_pages.append(pages[i].page_number)

    result = _assemble(
        p.name, cache_key, pages, tables, page_count, extraction_chain,
        [], ocr_pages, vision_pages, warnings, errors,
    )
    _extraction_cache[f"deep:{cache_key}"] = result
    return result


# ── Main entry point ──────────────────────────────────────────────────────────

def extract(
    path: str,
    provider=None,
    mode: str = "fast",
    ocr_semaphore=None,
) -> ParsedDocument:
    """
    Main extraction entry point.
    mode="fast"  → extract_fast (no OCR/vision unless critically weak)
    mode="deep"  → extract_deep (OCR + vision on weak pages)
    Non-PDF files always use direct extraction.
    ocr_semaphore: threading.Semaphore passed from router for OCR throttling.
    """
    p = Path(path)
    if p.suffix.lower() not in (".pdf",):
        return _extract_direct(path)

    if mode == "deep":
        return extract_deep(path, provider=provider, ocr_semaphore=ocr_semaphore)
    return extract_fast(path, provider=provider, ocr_semaphore=ocr_semaphore)
